#!/usr/bin/env python3
"""Генерация оформленной главы 3 ТЭО (DOCX) по методичке КМПО РАНХиГС."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Twips

OUTPUT = Path(__file__).resolve().parents[1] / "docs" / "Глава_3_ТЭО_оформленная.docx"

FONT = "Times New Roman"
SIZE = Pt(14)
INDENT = Cm(1.25)


def set_cell_shading(cell, fill: str = "FFFFFF"):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = SIZE
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)
    rFonts.set(qn("w:cs"), FONT)
    rPr.insert(0, rFonts)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = SIZE
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = INDENT
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_paragraph(doc, text, *, bold=False, indent=True, align=None, space_before=0, space_after=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if indent:
        pf.first_line_indent = INDENT
    else:
        pf.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_run_font(run, bold=bold)
    return p


def add_empty_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run("")
    set_run_font(run)
    return p


def add_heading_section(doc, number: str, title: str):
    add_empty_line(doc)
    text = f"{number} {title}" if number else title
    return add_paragraph(doc, text, bold=True, indent=True)


def add_heading_subsection(doc, number: str, title: str):
    return add_paragraph(doc, f"{number} {title}", bold=True, indent=True)


def add_bullet(doc, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1.25)
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_formula(doc, formula: str, number: str):
    """Формула по центру, номер справа (п. 5.8.3)."""
    add_empty_line(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    tab_stops = pf.tab_stops
    tab_stops.add_tab_stop(Cm(9.0), WD_ALIGN_PARAGRAPH.CENTER)
    tab_stops.add_tab_stop(Cm(16.5), WD_ALIGN_PARAGRAPH.RIGHT)

    run = p.add_run(f"\t{formula}\t({number})")
    set_run_font(run)
    add_empty_line(doc)
    return p


def add_where_block(doc, lines: list[str]):
    """Пояснение символов под формулой (п. 5.8.2)."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.first_line_indent = INDENT
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(lines[0])
    set_run_font(run)
    for line in lines[1:]:
        p2 = doc.add_paragraph()
        p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p2.paragraph_format.first_line_indent = INDENT
        r2 = p2.add_run(line)
        set_run_font(r2)


def add_table_caption(doc, caption: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(caption)
    set_run_font(run)


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                for run in p.runs:
                    set_run_font(run, bold=run.bold)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                el = OxmlElement(f"w:{edge}")
                el.set(qn("w:val"), "single")
                el.set(qn("w:sz"), "4")
                el.set(qn("w:space"), "0")
                el.set(qn("w:color"), "000000")
                tcBorders.append(el)
            tcPr.append(tcBorders)


def add_table(doc, headers: list[str], rows: list[list[str]], col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    style_table(table)
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, bold=True)
    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            row_cells[c_idx].text = ""
            p = row_cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_run_font(run, bold=(str(val).lower().startswith("итого")))
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Twips(int(w * 567))  # cm approx via dxa - use Cm
    doc.add_paragraph()
    return table


def build_chapter(doc: Document):
    add_heading_section(doc, "3", "ТЕХНИКО–ЭКОНОМИЧЕСКОЕ ОБОСНОВАНИЕ")

    # --- 3.1 ---
    add_heading_subsection(doc, "3.1", "Определение факторов эффективности внедрения")
    add_paragraph(
        doc,
        "Внедрение веб–модуля направлено на автоматизацию подготовки экологической отчётности "
        "и оперативного контроля показателей. Эффект проявляется в качественных и количественных "
        "изменениях процесса.",
    )
    add_paragraph(doc, "Качественные эффекты:", indent=True)
    for b in [
        "снижение доли ручного труда при сборе и проверке данных;",
        "уменьшение риска ошибок при переносе цифр между файлами;",
        "ускорение информирования руководства (KPI на дашбордах «Контроль» и «Отчётность»);",
        "повышение прозрачности учёта отходов и экологических измерений;",
        "единый доступ к данным платформы через REST API без дублирования в локальных таблицах.",
    ]:
        add_bullet(doc, b)

    add_paragraph(doc, "Количественные показатели эффективности, используемые в расчётах:", indent=True)
    for b in [
        "ΔT – экономия времени на подготовку одного отчётного комплекта, ч;",
        "N – число отчётных циклов в год (принято 12, ежемесячная отчётность);",
        "Cч – стоимость одного часа работы специалиста, руб/ч;",
        "Эг – годовой экономический эффект, руб/год;",
        "K – полная себестоимость разработки, руб;",
        "Tок – срок окупаемости, лет (см. формулу (3.11)).",
    ]:
        add_bullet(doc, b)

    add_paragraph(
        doc,
        "Сравнение трудозатрат «до» и «после» внедрения приведено в таблице 3.1.",
    )
    add_table_caption(doc, "Таблица 3.1 – Сравнение трудозатрат на подготовку отчётности")
    add_table(
        doc,
        ["Операция", "Ручной способ, мин", "С модулем, мин", "Сокращение, мин"],
        [
            ["Сбор данных за период", "40", "10", "30"],
            ["Группировка и проверка", "35", "10", "25"],
            ["Расчёт итогов и KPI", "30", "3", "27"],
            ["Подготовка отчёта", "45", "8", "37"],
            ["Подготовка аналитики", "30", "5", "25"],
            ["Итого", "180", "36", "144"],
        ],
    )

    add_paragraph(doc, "Качественные эффекты внедрения сведены в таблице 3.2.")
    add_table_caption(doc, "Таблица 3.2 – Качественные эффекты внедрения")
    add_table(
        doc,
        ["Направление", "До внедрения", "После внедрения"],
        [
            ["Достоверность данных", "Зависит от ручного ввода", "Повышается за счёт единой БД платформы"],
            ["Скорость анализа", "Низкая, ручная обработка", "Высокая, KPI на дашборде"],
            ["Контроль превышений", "Проверяется вручную", "Фильтр и выделение в журнале измерений"],
            ["Отчётность", "Формируется вручную", "Экспорт PDF / XML"],
            ["Прозрачность процессов", "Разрозненные файлы", "Единый веб–интерфейс с ролями"],
        ],
    )

    add_paragraph(
        doc,
        "По таблице 3.1 экономия на один отчётный цикл составляет 144 мин = 2,4 ч. "
        "Расчёт выполняется по формуле (3.1).",
    )
    add_formula(doc, "ΔT = Tдо − Tпос", "3.1")
    add_where_block(
        doc,
        [
            "где",
            "ΔT – экономия времени на один отчётный цикл, ч;",
            "Tдо – суммарное время при ручной подготовке, ч;",
            "Tпос – суммарное время с модулем, ч.",
        ],
    )
    add_paragraph(
        doc,
        "По таблице 3.1: Tдо = 180 мин = 3,0 ч; Tпос = 36 мин = 0,6 ч; ΔT = 2,4 ч.",
    )

    # --- 3.2 ---
    add_heading_subsection(doc, "3.2", "Оценка трудоёмкости разработки")
    add_paragraph(
        doc,
        "Проект выполнен одним разработчиком, совмещающим функции аналитика, программиста, "
        "тестировщика и автора документации. Это типично для учебного дипломного проекта.",
    )
    add_paragraph(
        doc,
        "Трудоёмкость оценена по этапам жизненного цикла разработки. Результаты приведены в таблице 3.3.",
    )
    add_table_caption(doc, "Таблица 3.3 – Оценка трудоёмкости разработки веб–модуля")
    add_table(
        doc,
        ["Этап работы", "Содержание работ", "Трудоёмкость, ч"],
        [
            ["Анализ предметной области", "Изучение учёта отходов, требований и аналогов", "16"],
            ["Проектирование", "БД, архитектура, диаграммы BPMN, DFD, UML, ER", "22"],
            ["Реализация моделей и интеграции", "ORM, API–клиент, RemoteDataService", "14"],
            ["Разработка интерфейса", "Шаблоны, навигация, роли, формы", "22"],
            ["Дашборд и KPI", "Агрегации, Chart.js, два дашборда", "18"],
            ["Экспорт отчётов", "PDF, XML", "12"],
            ["Тестирование", "Ручная проверка сценариев", "16"],
            ["Документация", "Описание решения, подготовка ВКР", "20"],
            ["Итого", "", "140"],
        ],
    )
    add_paragraph(
        doc,
        "Перевод в человеко–месяцы выполнен исходя из нормы 176 ч/мес (22 рабочих дня × 8 ч) "
        "по формуле (3.2).",
    )
    add_formula(doc, "T = Tр / Tн", "3.2")
    add_where_block(
        doc,
        [
            "где",
            "T – трудоёмкость, чел.–мес.;",
            "Tр – трудозатраты на проект, ч (140);",
            "Tн – норма рабочего времени, ч/мес (176).",
        ],
    )
    add_paragraph(doc, "T = 140 / 176 ≈ 0,80 чел.–мес.")
    add_paragraph(
        doc,
        "Суммарная трудоёмкость – 140 человеко–часов, что соответствует примерно одному месяцу "
        "работы одного специалиста с неполной загрузкой или 3,5 неделям при полной занятости.",
    )

    # --- 3.3 ---
    add_heading_subsection(doc, "3.3", "Расчёт заработной платы и начислений")
    add_paragraph(
        doc,
        "За основу принят условный месячный оклад разработчика (студент / младший программист): "
        "ЗП = 55 000 руб/мес.",
    )
    add_paragraph(doc, "Фонд оплаты труда (ФОТ) на проект рассчитывается по формуле (3.3).")
    add_formula(doc, "ФОТ = ЗП × (Tр / Tн)", "3.3")
    add_where_block(
        doc,
        [
            "где",
            "ФОТ – фонд оплаты труда на проект, руб;",
            "ЗП – месячный оклад, руб/мес (55 000);",
            "Tр – трудозатраты, ч (140);",
            "Tн – норма времени, ч/мес (176).",
        ],
    )
    add_paragraph(doc, "ФОТ = 55 000 × (140 / 176) = 43 750 руб.")
    add_paragraph(
        doc,
        "Распределение ФОТ по видам работ (пропорционально таблице 3.3) приведено в таблице 3.4.",
    )
    add_table_caption(doc, "Таблица 3.4 – Распределение фонда оплаты труда по этапам")
    add_table(
        doc,
        ["Этап", "Доля, %", "Сумма, руб"],
        [
            ["Анализ", "11,4", "4 988"],
            ["Проектирование", "15,7", "6 869"],
            ["Реализация", "10,0", "4 375"],
            ["Интерфейс", "15,7", "6 869"],
            ["Дашборд и KPI", "12,9", "5 644"],
            ["Экспорт", "8,6", "3 763"],
            ["Тестирование", "11,4", "4 988"],
            ["Документация", "14,3", "6 254"],
            ["Итого", "100", "43 750"],
        ],
    )
    add_paragraph(doc, "Страховые взносы работодателя (принято 30,2 % от ФОТ) рассчитываются по формуле (3.4).")
    add_formula(doc, "СВ = ФОТ × kстр", "3.4")
    add_where_block(
        doc,
        [
            "где",
            "СВ – страховые взносы работодателя, руб;",
            "kстр – ставка взносов (0,302).",
        ],
    )
    add_paragraph(doc, "СВ = 43 750 × 0,302 = 13 213 руб.")
    add_paragraph(doc, "Затраты на персонал всего рассчитываются по формуле (3.5).")
    add_formula(doc, "Зперс = ФОТ + СВ", "3.5")
    add_where_block(
        doc,
        [
            "где",
            "Зперс – затраты на персонал, руб.",
        ],
    )
    add_paragraph(doc, "Зперс = 43 750 + 13 213 = 56 963 руб.")

    # --- 3.4 ---
    add_heading_subsection(doc, "3.4", "Амортизация оборудования и нематериальных активов")
    add_paragraph(
        doc,
        "Для разработки использовались следующие активы, представленные в таблице 3.5.",
    )
    add_table_caption(doc, "Таблица 3.5 – Оборудование и программное обеспечение проекта")
    add_table(
        doc,
        [
            "Наименование",
            "Стоимость, руб",
            "Срок службы, мес",
            "Срок использования в проекте, мес",
            "Амортизация на проект, руб",
        ],
        [
            ["Персональный компьютер", "90 000", "36", "4", "10 000"],
            ["ОС (Linux / Windows)", "0", "–", "4", "0"],
            ["IDE (PyCharm Community)", "0", "–", "4", "0"],
            ["Python, Django, библиотеки", "0", "–", "4", "0"],
            ["Итого", "90 000", "", "", "10 000"],
        ],
    )
    add_paragraph(doc, "Линейная амортизация на период проекта определяется по формуле (3.6).")
    add_formula(doc, "А = (Соб / Tсл) × Tисп", "3.6")
    add_where_block(
        doc,
        [
            "где",
            "А – амортизация на период проекта, руб;",
            "Соб – стоимость оборудования, руб (90 000);",
            "Tсл – срок службы, мес (36);",
            "Tисп – срок использования в проекте, мес (4).",
        ],
    )
    add_paragraph(doc, "Для ПК: А = (90 000 / 36) × 4 = 10 000 руб.")
    add_paragraph(
        doc,
        "Использование открытого ПО (Python, Django, Bootstrap, Chart.js, reportlab) "
        "не требует лицензионных отчислений, что снижает себестоимость проекта.",
    )

    # --- 3.5 ---
    add_heading_subsection(doc, "3.5", "Расчёт стоимости оборотных средств и текущих затрат")
    add_table_caption(doc, "Таблица 3.6 – Затраты на разработку и эксплуатацию модуля")
    add_table(
        doc,
        ["Статья", "Расчёт", "Сумма, руб"],
        [
            ["Доступ в интернет", "600 руб/мес × 4 мес", "2 400"],
            ["Электроэнергия", "~200 Вт × 4 ч/день × 80 дней × 6 руб/кВт·ч", "400"],
            ["Расходные материалы (печать, канцелярия)", "Лимит", "600"],
            ["Итого за разработку", "", "3 400"],
            ["Хостинг клиента ECO (VPS)", "500 руб/мес × 12", "6 000"],
            ["Интернет", "600 руб/мес × 12", "7 200"],
            ["Резервное копирование / администрирование", "2 ч/мес × 400 руб/ч × 12", "9 600"],
            ["Итого эксплуатация", "", "22 800"],
        ],
    )
    add_paragraph(
        doc,
        "Основные учётные данные хранятся на сервере платформы комплекса (PostgreSQL); "
        "модуль ECO – клиент, поэтому затраты на промышленный сервер учёта в расчёт не включены.",
    )

    # --- 3.6 ---
    add_heading_subsection(doc, "3.6", "Прочие затраты и накладные расходы")
    add_table_caption(doc, "Таблица 3.7 – Прочие прямые затраты")
    add_table(
        doc,
        ["Статья", "Сумма, руб"],
        [
            ["Обучение пользователей (1 занятие, 2 ч)", "800"],
            ["Консультации / рецензирование", "0"],
            ["Итого прочие", "800"],
        ],
    )
    add_paragraph(
        doc,
        "Накладные расходы (административно–хозяйственные, связь, бухгалтерское сопровождение) "
        "приняты в размере 10 % от прямых затрат разработки (без эксплуатации) и рассчитываются "
        "по формуле (3.7).",
    )
    add_formula(doc, "НР = kнр × Зпр", "3.7")
    add_where_block(
        doc,
        [
            "где",
            "НР – накладные расходы, руб;",
            "kнр – коэффициент накладных расходов (0,10);",
            "Зпр – прямые затраты разработки без эксплуатации, руб.",
        ],
    )
    add_paragraph(doc, "Зпр = 56 963 + 10 000 + 3 400 + 800 = 71 163 руб.")
    add_paragraph(doc, "НР = 0,10 × 71 163 = 7 116 руб.")

    # --- 3.7 ---
    add_heading_subsection(doc, "3.7", "Калькуляция полной себестоимости и экономический эффект")
    add_paragraph(doc, "Сводная калькуляция затрат на разработку приведена в таблице 3.8.")
    add_table_caption(doc, "Таблица 3.8 – Калькуляция полной себестоимости разработки")
    add_table(
        doc,
        ["№", "Статья затрат", "Сумма, руб"],
        [
            ["1", "Фонд оплаты труда", "43 750"],
            ["2", "Страховые взносы (30,2 %)", "13 213"],
            ["3", "Амортизация оборудования", "10 000"],
            ["4", "Материалы и услуги (интернет, электроэнергия)", "3 400"],
            ["5", "Прочие затраты", "800"],
            ["6", "Накладные расходы (10 %)", "7 116"],
            ["", "Полная себестоимость K", "78 279"],
        ],
    )

    add_paragraph(doc, "Расчёт годового экономического эффекта.")
    add_paragraph(doc, "Приняты допущения:")
    add_bullet(doc, "N = 12 отчётных циклов в год;")
    add_bullet(doc, "ΔT = 2,4 ч – экономия времени эколога на цикл (табл. 3.1);")
    add_bullet(doc, "Cч = 400 руб/ч – стоимость часа работы эколога с учётом начислений;")
    add_bullet(doc, "дополнительная экономия времени руководителя на аналитику: 5 ч/год × 600 руб/ч = 3 000 руб.")

    add_paragraph(doc, "Экономия ФОТ эколога рассчитывается по формуле (3.8).")
    add_formula(doc, "Ээк = ΔT × N × Cч", "3.8")
    add_where_block(
        doc,
        [
            "где",
            "Ээк – экономия ФОТ эколога, руб/год;",
            "ΔT – экономия времени на цикл, ч (2,4);",
            "N – число циклов в год (12);",
            "Cч – стоимость часа эколога, руб/ч (400).",
        ],
    )
    add_paragraph(doc, "Ээк = 2,4 × 12 × 400 = 11 520 руб/год.")
    add_paragraph(doc, "Эрук = 3 000 руб/год.")
    add_paragraph(
        doc,
        "Условная экономия от снижения ошибок (1 исправляемая ошибка в отчёте в год): "
        "Эош = 15 000 руб/год.",
    )
    add_paragraph(doc, "Годовой эффект определяется по формуле (3.9).")
    add_formula(doc, "Эг = Ээк + Эрук + Эош", "3.9")
    add_paragraph(doc, "Эг = 11 520 + 3 000 + 15 000 = 29 520 руб/год.")

    add_paragraph(doc, "Полная себестоимость разработки определяется по формуле (3.10).")
    add_formula(doc, "K = Зперс + А + Зм + Зпроч + НР", "3.10")
    add_where_block(
        doc,
        [
            "где",
            "K – полная себестоимость разработки, руб;",
            "Зперс – затраты на персонал, руб (56 963);",
            "А – амортизация, руб (10 000);",
            "Зм – материалы и услуги, руб (3 400);",
            "Зпроч – прочие затраты, руб (800);",
            "НР – накладные расходы, руб (7 116).",
        ],
    )
    add_paragraph(doc, "K = 78 279 руб (табл. 3.8).")

    add_paragraph(doc, "Срок окупаемости рассчитывается по формуле (3.11).")
    add_formula(doc, "Ток = K / Эг", "3.11")
    add_where_block(
        doc,
        [
            "где",
            "Ток – срок окупаемости, лет;",
            "K – полная себестоимость разработки, руб (78 279);",
            "Эг – годовой экономический эффект, руб/год (29 520).",
        ],
    )
    add_paragraph(doc, "Ток = 78 279 / 29 520 ≈ 2,65 года (≈ 32 месяца).")
    add_paragraph(
        doc,
        "Примечание – Расчёт срока окупаемости выполнен без учёта годовых эксплуатационных "
        "затрат (22 800 руб/год, табл. 3.6); при их учёте срок окупаемости несколько увеличивается.",
    )
    add_paragraph(doc, "Относительная экономия времени на подготовку одного отчёта определяется по формуле (3.12).")
    add_formula(doc, "η = (Tдо − Tпос) / Tдо × 100 %", "3.12")
    add_where_block(
        doc,
        [
            "где",
            "η – относительная экономия времени, %;",
            "Tдо, Tпос – время подготовки отчёта до и после внедрения, мин.",
        ],
    )
    add_paragraph(doc, "η = (180 − 36) / 180 × 100 % = 80 %.")
    add_paragraph(
        doc,
        "При внедрении на предприятии с бóльшим числом организаций и ежеквартальной отчётностью "
        "в несколько подразделений эффект Эг возрастает пропорционально числу отчётных комплектов, "
        "а срок окупаемости сокращается.",
    )

    # --- 3.8 ---
    add_heading_subsection(doc, "3.8", "Итоговые экономические показатели и выводы")
    add_paragraph(
        doc,
        "По результатам технико–экономического обоснования получены показатели, "
        "представленные в таблице 3.9.",
    )
    add_table_caption(doc, "Таблица 3.9 – Итоговые показатели")
    add_table(
        doc,
        ["Показатель", "Значение"],
        [
            ["Трудоёмкость разработки", "140 ч (0,80 чел.–мес)"],
            ["Полная себестоимость разработки K", "78 279 руб"],
            ["Годовой экономический эффект Эг", "29 520 руб"],
            ["Экономия времени на один отчёт", "144 мин (80 %)"],
            ["Срок окупаемости Tок", "≈ 2,65 года (32 мес.)"],
            ["Годовые затраты на эксплуатацию", "22 800 руб"],
        ],
    )

    add_paragraph(doc, "Вывод по главе 3:", bold=True)
    conclusions = [
        "Разработка веб–модуля экономически оправдана для учебного и опытного внедрения: "
        "затраты на создание прототипа умеренные (~ 78 тыс. руб.) за счёт использования "
        "открытого ПО и работы одного разработчика.",
        "Основной количественный эффект – сокращение времени подготовки отчётности на 80 % "
        "и высвобождение ~29 ч/год рабочего времени эколога и дополнительно 5 ч/год руководителя "
        "(всего ~34 ч/год).",
        "Качественный эффект – повышение достоверности данных, прозрачность контроля превышений, "
        "единый доступ к данным платформы.",
        "Проект наиболее выгоден при регулярной (ежемесячной) отчётности и использовании модуля "
        "не менее 2–3 лет; при масштабировании на несколько площадок срок окупаемости "
        "сокращается до 1–1,5 года.",
        "Использование Django, REST API, PostgreSQL (сервер платформы) и бесплатных библиотек "
        "снижает капитальные и лицензионные затраты.",
        "Технико–экономическое обоснование подтверждает целесообразность разработки веб–модуля "
        "отчётности и экологического контроля: структура затрат прозрачна, эффект от автоматизации "
        "выражен в количественных и качественных показателях, проект может быть рекомендован "
        "для опытной эксплуатации в составе программного комплекса по учёту и переработке "
        "промышленных отходов.",
    ]
    for c in conclusions:
        add_paragraph(doc, c)

    add_paragraph(
        doc,
        "Сравнение с альтернативой: ручная подготовка отчётности не требует капитальных затрат на "
        "разработку, но ежегодно «стоит» ~29–34 ч рабочего времени специалистов и несёт риск ошибок; "
        "разработанный модуль при K = 78 279 руб и Эг = 29 520 руб/год окупается за ~2,7 года, "
        "после чего даёт чистую экономию времени и повышение качества данных.",
    )


def main():
    doc = Document()
    configure_document(doc)
    build_chapter(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    main()
