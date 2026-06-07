#!/usr/bin/env python3
"""Insert figure placeholders and captions into diploma docx."""

from __future__ import annotations

import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

DOCX = Path(__file__).resolve().parents[2] / "16_Долгов ЕВ_01062026.docx"


def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def insert_caption_after(paragraph: Paragraph, caption: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if " – " in caption:
        label, rest = caption.split(" – ", 1)
        run = new_para.add_run(label)
        run.bold = True
        new_para.add_run(" – " + rest)
    else:
        new_para.add_run(caption)
    return new_para


def insert_placeholder_after(paragraph: Paragraph, text: str) -> Paragraph:
    para = insert_paragraph_after(paragraph, text)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return para


def insert_figure_block(
    anchor: Paragraph,
    *,
    intro: str | None = None,
    placeholder: str,
    caption: str,
    after: str | None = None,
) -> Paragraph:
    current = anchor
    if intro:
        current = insert_paragraph_after(current, intro)
    current = insert_placeholder_after(current, placeholder)
    current = insert_caption_after(current, caption)
    if after:
        current = insert_paragraph_after(current, after)
    return current


def find_para(doc: Document, needle: str) -> Paragraph:
    for p in doc.paragraphs:
        if needle in p.text:
            return p
    raise ValueError(f"Paragraph not found: {needle!r}")


def main() -> None:
    if not DOCX.exists():
        print(f"File not found: {DOCX}", file=sys.stderr)
        sys.exit(1)

    backup = DOCX.with_suffix(".docx.bak")
    shutil.copy2(DOCX, backup)
    print(f"Backup: {backup}")

    doc = Document(DOCX)

    # --- Figure 10 → 9 → … → 1 (bottom-up by anchor uniqueness) ---

    p = find_para(doc, "Для промышленной эксплуатации предусматриваются фильтры")
    insert_figure_block(
        p,
        intro=(
            "Центр уведомлений отображает регламентные предупреждения T1 (прогноз простоя на линии) "
            "и T2 (риск нарушения срока хранения партии). Экран прототипа показан на рисунке 10."
        ),
        placeholder="[ВСТАВИТЬ СКРИН — вкладка «Уведомления»]",
        caption="Рисунок 10 – Экран «Центр уведомлений» веб-модуля планирования (триггеры T1/T2)",
    )

    p = find_para(doc, "Интерфейс организуется вокруг рабочего места диспетчера")
    insert_figure_block(
        p,
        intro=(
            "Сценарное моделирование «что если» (базовый, ускоренный, аварийный) и сравнение KPI "
            "двух вариантов плана реализованы на отдельной вкладке интерфейса (рисунок 9)."
        ),
        placeholder="[ВСТАВИТЬ СКРИН — вкладка «Симуляция» + таблица результата]",
        caption="Рисунок 9 – Экран сценарного моделирования и сравнения KPI базового и симулированного плана",
    )

    p = find_para(doc, "Экран «Планирование смены» включает панель действий")
    insert_figure_block(
        p,
        intro=(
            "Экран планирования смены включает Gantt-диаграмму загрузки линий, таблицу операций "
            "и действия диспетчера (построение, утверждение, перепланирование). "
            "Интерфейс прототипа представлен на рисунке 8."
        ),
        placeholder="[ВСТАВИТЬ СКРИН — вкладка «Расписание»]",
        caption="Рисунок 8 – Экран «Расписание переработки»: Gantt-диаграмма и таблица операций плана",
        after=(
            "В прототипе обновление данных выполняется по запросу пользователя (кнопка «Обновить»); "
            "маркер текущего времени на шкале Gantt и push-обновление метрик запланированы в следующей версии."
        ),
    )

    p = find_para(doc, "Графическое представление пользовательского интерфейса приведено в приложении Г")
    p.text = (
        "Графическое представление пользовательского интерфейса приведено в приложении Г "
        "(макеты экрана планирования смены с диаграммой Ганта и оперативного дашборда). "
        "Проектирование выполнено с учётом ГОСТ Р ИСО 9241-210: приоритет отдан задачам диспетчера — "
        "построение и публикация плана, контроль загрузки линий, реагирование на риски."
    )
    insert_figure_block(
        p,
        intro=(
            "Экран оперативного обзора с ключевыми показателями эффективности (OEE, выполнение плана, "
            "простой, риск хранения) и загрузкой производственных линий реализован в веб-клиенте "
            "и показан на рисунке 7."
        ),
        placeholder="[ВСТАВИТЬ СКРИН — вкладка «Обзор»]",
        caption="Рисунок 7 – Экран «Обзор производства» веб-модуля планирования: KPI и загрузка линий L1/L2",
    )

    p = find_para(doc, "Логическая структура данных модуля отражена ER-диаграммой")
    p.text = (
        "Логическая структура данных модуля отражена ER-диаграммой (приложение Б, рисунок 6). "
        "Модуль использует единую реляционную базу данных программного комплекса (PostgreSQL); "
        "имена таблиц и полей приведены в регистре snake_case, типы атрибутов — в терминах целевой СУБД."
    )
    insert_figure_block(
        p,
        intro="Логическая модель данных модуля планирования приведена на рисунке 6.",
        placeholder="[ВСТАВИТЬ РИСУНОК 6]",
        caption=(
            "Рисунок 6 – ER-диаграмма базы данных модуля планирования "
            "(фрагмент общей СУБД комплекса)"
        ),
    )

    p = find_para(doc, "Структура программного обеспечения представлена UML-диаграммой компонентов")
    p.text = (
        "Структура программного обеспечения представлена UML-диаграммой компонентов "
        "(приложение В, рисунок 5). Модуль реализован как тонкий клиент: веб-приложение на React "
        "и TypeScript обращается к API программного комплекса, не имея собственной базы данных "
        "и не дублируя бизнес-логику планирования."
    )
    insert_figure_block(
        p,
        intro=(
            "Структура программного обеспечения модуля планирования представлена "
            "UML-диаграммой компонентов (рисунок 5, приложение В)."
        ),
        placeholder="[ВСТАВИТЬ РИСУНОК 5]",
        caption=(
            "Рисунок 5 – UML-диаграмма компонентов веб-модуля планирования перерабатывающих процессов"
        ),
        after=(
            "В реализованном прототипе веб-клиент обращается к API платформы по REST (HTTPS/JSON); "
            "бизнес-логика планирования, симуляции и уведомлений выполняется на сервере FastAPI "
            "с доступом к PostgreSQL. Обмен событиями в режиме push (SSE) и интеграция с MES/SCADA "
            "в прототипе не реализованы и отнесены к перспективе развития."
        ),
    )

    p = find_para(doc, "Данные, передаваемые модулю мониторинга")
    insert_figure_block(
        p,
        intro=(
            "Передача плановых данных смежным модулям после утверждения версии расписания "
            "показана на рисунке 4."
        ),
        placeholder="[ВСТАВИТЬ РИСУНОК 4]",
        caption=(
            "Рисунок 4 – Передача данных утверждённого плана модулям мониторинга и отчётности"
        ),
    )

    p = find_para(doc, "Обмен между пулами диспетчера, модуля и производства моделируется")
    p.text = p.text.replace("на рисунке)", "на рисунке 3)")
    insert_figure_block(
        p,
        placeholder="[ВСТАВИТЬ РИСУНОК 3]",
        caption="Рисунок 3 – BPMN-модель процесса планирования переработки промышленных отходов",
        after=(
            "На рисунке 3 выделены пул диспетчера (построение и утверждение плана), "
            "пул модуля автоматического планирования (расчёт приоритетов, расписания, "
            "симуляция при конфликтах) и пул производства (исполнение утверждённого плана)."
        ),
    )

    p = find_para(doc, "Формализованное описание сменного цикла")
    p.text = (
        "Формализованное описание сменного цикла «ввод данных — расчёт — утверждение — исполнение» "
        "выполнено в нотации BPMN и представлено на рисунке 3 (полный вариант — приложение А). "
        "На диаграмме выделены три пула ответственности: «Диспетчер», «Модуль планирования» "
        "и «Производство / Технолог», что позволяет разграничить действия пользователя "
        "и автоматизированной обработки."
    )

    p = find_para(doc, "Продолжение таблицы 2")
    insert_figure_block(
        p,
        intro=(
            "Сводное сопоставление текущего (as-is) и целевого (to-be) процесса планирования "
            "переработки отходов представлено на рисунке 2."
        ),
        placeholder="[ВСТАВИТЬ РИСУНОК 2]",
        caption=(
            "Рисунок 2 – Сравнение процесса планирования переработки отходов: "
            "текущее и целевое состояние"
        ),
    )

    p = find_para(doc, "Продолжение таблицы 1")
    insert_figure_block(
        p,
        intro=(
            "Разрабатываемый веб-модуль планирования встроен в единый программный комплекс учёта "
            "и переработки промышленных отходов и обменивается данными с модулями учёта, мониторинга "
            "и отчётности через общую СУБД PostgreSQL и REST API (рисунок 1)."
        ),
        placeholder="[ВСТАВИТЬ РИСУНОК 1]",
        caption=(
            "Рисунок 1 – Место модуля планирования перерабатывающих процессов в программном комплексе"
        ),
    )

    # --- Appendices ---
    p = find_para(doc, "Приложение А. BPMN-диаграмма процессов планирования")
    insert_placeholder_after(p, "[ВСТАВИТЬ ДИАГРАММУ — BPMN (полный лист, см. рисунок 3)]")

    p = find_para(doc, "Приложение Б. ER-диаграмма базы данных модуля планирования")
    insert_placeholder_after(p, "[ВСТАВИТЬ ДИАГРАММУ — ER-модель (полный лист, см. рисунок 6)]")

    p = find_para(doc, "Приложение В. Диаграмма прецендентов")
    p.text = "Приложение В. UML-диаграмма компонентов/архитектуры"
    insert_placeholder_after(p, "[ВСТАВИТЬ ДИАГРАММУ — UML компонентов (полный лист, см. рисунок 5)]")

    p = find_para(doc, "Приложение Г. Макеты интерфейса")
    p.text = (
        "Приложение Г. Макеты интерфейса веб-модуля планирования "
        "(лист 1 — обзор; лист 2 — расписание; лист 3 — симуляция; лист 4 — уведомления)"
    )
    insert_placeholder_after(
        p,
        "[ВСТАВИТЬ СКРИНЫ — лист 1: обзор (рис. 7); лист 2: расписание (рис. 8); "
        "лист 3: симуляция (рис. 9); лист 4: уведомления (рис. 10)]",
    )

    # Remove duplicate appendix Д (UML moved to В)
    for p in doc.paragraphs:
        if p.text.strip() == "Приложение Д. Диаграмма компонентов":
            el = p._element
            el.getparent().remove(el)
            break

    doc.save(DOCX)
    print(f"Updated: {DOCX}")


if __name__ == "__main__":
    main()
